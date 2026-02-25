from __future__ import annotations

import io
import json
import os
import re
import shutil
import subprocess
import tempfile
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple

from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate

from app.services.llm_router import invoke_with_fallback


_CHARTS_START = "===CHARTS_JSON==="
_CHARTS_END = "===END_CHARTS_JSON==="


def _format_docs_for_report(docs: List[Document]) -> str:
    parts: list[str] = []
    for doc in docs:
        source = doc.metadata.get("source_file", "Unknown Source")
        page = doc.metadata.get("page", None)
        parts.append(
            "Content:\n"
            f"{doc.page_content}\n\n"
            f"Source File: {source}\n"
            f"Page: {page}\n"
        )
    return "\n\n---\n\n".join(parts)


def _build_report_prompt() -> ChatPromptTemplate:
    system = (
        'You are "NashmiBot Report Writer", an official government-grade analyst.\n'
        "Write a clear, well-structured report based ONLY on the provided context.\n\n"
        "STRICT RULES:\n"
        "1) Use ONLY the provided context. Do not invent facts.\n"
        "2) Every factual claim MUST include an in-text citation in this format:\n"
        "   [Source: filename.pdf p. 12]  (page can be None -> omit p.)\n"
        "3) If context is insufficient for the requested report, write what is supported and explicitly list missing data.\n"
        "4) Output in the SAME language as the user's topic (Arabic or English).\n"
        "5) Use simple Markdown structure: #, ##, bullet points (- ), and short paragraphs.\n\n"
        "CHARTS (optional):\n"
        "If (and only if) the context contains clear numeric series suitable for visualization,\n"
        "append a charts JSON block using EXACT markers:\n"
        f"{_CHARTS_START}\n"
        # Escape braces for ChatPromptTemplate formatting.
        '[{{"type":"bar|line","title":"...","x":["..."],"y":[1,2,3],"x_label":"...","y_label":"...","citations":["[Source: ...]"]}}]\n'
        f"{_CHARTS_END}\n"
        "If no charts are possible, output exactly an empty list in that block.\n\n"
        "Context:\n{context}\n"
    )

    return ChatPromptTemplate.from_messages(
        [
            ("system", system),
            ("human", "{topic}"),
        ]
    )


def _extract_charts_block(text: str) -> Tuple[str, List[Dict[str, Any]]]:
    if _CHARTS_START not in text or _CHARTS_END not in text:
        return text.strip(), []

    before, rest = text.split(_CHARTS_START, 1)
    block, after = rest.split(_CHARTS_END, 1)
    md = (before + after).strip()

    raw = block.strip()
    try:
        data = json.loads(raw)
        if isinstance(data, list):
            return md, [d for d in data if isinstance(d, dict)]
    except Exception:
        pass

    return md, []


def generate_report_markdown(topic: str, docs: List[Document], provider_preference: str = "auto") -> Tuple[str, List[Dict[str, Any]]]:
    context = _format_docs_for_report(docs)
    prompt = _build_report_prompt()
    text, _provider = invoke_with_fallback(
        prompt,
        {"context": context, "topic": topic},
        # Reports can be longer than chat answers.
        max_output_tokens=3072,
        temperature=0.0,
        provider_preference=provider_preference,
    )
    md, charts = _extract_charts_block(text)
    return md, charts


def _safe_filename_slug(text: str, max_len: int = 50) -> str:
    # Keep filenames ASCII-only to avoid header encoding issues across proxies/browsers.
    text = text.strip().lower()
    text = re.sub(r"[^a-z0-9\s-]", "", text)
    text = re.sub(r"\s+", "-", text).strip("-")
    return (text[:max_len] or "report").strip("-")


def build_docx_report(
    *,
    topic: str,
    markdown: str,
    charts: List[Dict[str, Any]],
    retrieved_docs: List[Document],
    include_charts: bool = True,
) -> bytes:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    # Lazy import: matplotlib can be heavy.
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    doc = DocxDocument()

    # Title
    doc.add_heading("NashmiBot Report", level=0)
    doc.add_paragraph(topic)
    doc.add_paragraph(datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"))

    doc.add_paragraph("")  # spacing

    # Minimal markdown -> docx rendering
    for raw_line in (markdown or "").splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            p.paragraph_format.space_after = 0
            continue

        doc.add_paragraph(line.strip())

    # Charts (optional)
    if include_charts and charts:
        valid_charts = 0
        for spec in charts[:6]:
            ctype = str(spec.get("type", "")).strip().lower()
            title = str(spec.get("title", "")).strip() or "Chart"
            x = spec.get("x")
            y = spec.get("y")
            if not isinstance(x, list) or not isinstance(y, list) or len(x) != len(y) or not x:
                continue
            try:
                y_vals = [float(v) for v in y]
            except Exception:
                continue

            fig = plt.figure(figsize=(8.0, 4.0), dpi=150)
            ax = fig.add_subplot(111)
            if ctype == "line":
                ax.plot(list(range(len(x))), y_vals, marker="o")
                ax.set_xticks(list(range(len(x))))
                ax.set_xticklabels([str(v) for v in x], rotation=30, ha="right")
            else:
                # default to bar
                ax.bar([str(v) for v in x], y_vals)
                ax.tick_params(axis="x", rotation=30)

            ax.set_title(title)
            ax.set_xlabel(str(spec.get("x_label", "") or ""))
            ax.set_ylabel(str(spec.get("y_label", "") or ""))
            fig.tight_layout()

            png = io.BytesIO()
            fig.savefig(png, format="png")
            plt.close(fig)
            png.seek(0)

            doc.add_paragraph("")
            doc.add_heading(title, level=2)
            doc.add_picture(png, width=Inches(6.5))

            citations = spec.get("citations")
            if isinstance(citations, list) and citations:
                doc.add_paragraph("Sources: " + " ".join(str(c) for c in citations))

            valid_charts += 1

        if valid_charts == 0:
            # Avoid empty section noise; no charts were valid.
            pass

    # Sources appendix
    sources: list[str] = []
    for d in retrieved_docs:
        src = d.metadata.get("source_file", None)
        page = d.metadata.get("page", None)
        if not src:
            continue
        label = f"{src} p.{page}" if page is not None else str(src)
        sources.append(label)
    sources = sorted(set(sources))
    if sources:
        doc.add_page_break()
        doc.add_heading("Sources Used", level=1)
        for s in sources:
            doc.add_paragraph(s, style="List Bullet")

    # Best-effort RTL alignment if the topic is Arabic.
    if re.search(r"[\u0600-\u06FF]", topic):
        for p in doc.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def convert_docx_bytes_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert DOCX bytes to PDF bytes (best-effort).

    Tries:
      1) docx2pdf (Windows / MS Word)
      2) LibreOffice 'soffice' (if available)
    """
    with tempfile.TemporaryDirectory() as td:
        in_path = f"{td}/report.docx"
        out_pdf = f"{td}/report.pdf"

        with open(in_path, "wb") as f:
            f.write(docx_bytes)

        # 1) docx2pdf
        try:
            from docx2pdf import convert as docx2pdf_convert

            docx2pdf_convert(in_path, out_pdf)
            with open(out_pdf, "rb") as f:
                return f.read()
        except Exception:
            pass

        # 2) LibreOffice
        soffice = shutil.which("soffice")
        if soffice:
            try:
                subprocess.run(
                    [soffice, "--headless", "--nologo", "--convert-to", "pdf", "--outdir", td, in_path],
                    check=True,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    timeout=60,
                )
                # LibreOffice names output as report.pdf based on input basename.
                if not os.path.exists(out_pdf):
                    # Some versions might produce upper-case extension or different basename.
                    candidates = [p for p in os.listdir(td) if p.lower().endswith(".pdf")]
                    if candidates:
                        out_pdf = os.path.join(td, candidates[0])
                with open(out_pdf, "rb") as f:
                    return f.read()
            except Exception as e:
                raise RuntimeError("Failed to convert DOCX to PDF via LibreOffice.") from e

        raise RuntimeError(
            "PDF export is unavailable on this server. "
            "Install Microsoft Word (docx2pdf) or LibreOffice (soffice) for DOCX->PDF conversion."
        )


def build_report_filename(topic: str, ext: str) -> str:
    slug = _safe_filename_slug(topic)
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M")
    return f"nashmibot-report-{slug}-{stamp}.{ext}"
